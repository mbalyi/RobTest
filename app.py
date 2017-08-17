from upload import UP
from dbinterface import DB
from messenger import Messenger
from reportinterface import Report
import json, os, base64, sys, linecache
from flask import Flask, render_template, session, redirect, url_for, escape, request, jsonify, Response, send_from_directory
from werkzeug.utils import secure_filename
#from flask_debugtoolbar import DebugToolbarExtension
from docx import Document
from docx.shared import Inches
import pdfkit
import xlsxwriter
import ast
#from PIL import Image

""" 
Necessary packages: 
	- nodejs
	- flask:		pip install flask
	- python-docx: 	pip install python-docx
	- pdfkit: 	pip install pdfkit
				pip install wkhtmltopdf
		Differencies between windows and linux. See details on:
		https://pypi.python.org/pypi/pdfkit/0.4.1
		On windows: 
					1.Downloads latest wkhtmltopdf installer.
					2.Copy the folder to the Tool folder.
					3.Set the absolute path in the functions
		Have to give difference absolute file path!
	- xlsxwriter:	pip install xlsxwriter
Client side:
	All js and css files attached! 
"""

UPLOAD_FOLDER = './uploads/'
ALLOWED_EXTENSIONS = set(['txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'doc', 'docx', 'csv','doc', 'docx', 'xlsx', 'xlt', 'xls', 'PNG', 'JPG', 'JPEG','html','htm','HTML','HTM'])

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/', methods=['GET', 'POST'])
def login(active=None):
	return render_template('login.html')

@app.route('/uploads/<path:filename>', methods=['GET', 'POST'])
def uploads(filename):
	if os.path.isdir(filename):
		return "dir"
	else:
		return send_from_directory("./uploads/", filename)
		
@app.route('/downloads/<path:filename>', methods=['GET', 'POST'])
def downloads(filename):
	return send_from_directory("./downloads/", filename)

@app.route('/templates/pdfTemplates/<path:filename>', methods=['GET', 'POST'])
def pdfTemplates(filename):
	return send_from_directory("./templates/pdfTemplates/", filename)
	
@app.route('/ROB_2016.s3db', methods=['GET', 'POST'])
def downloadBackup():
	return send_from_directory("./", "ROB_2016.s3db")
	
@app.route('/home', methods=['GET', 'POST'])
def login_form():
	log = DB.login_querry(title=request.form["user"],pw=request.form["password"])
	session['username'] = request.form['user']
	if log == 1 and 'username' in session:
		projects = DB.getProjects()
		projectId = DB.getSelectedProject(user=request.form['user'])
		admin=DB.isAdmin(user=request.form['user'])
		return render_template('index.html',user = admin, Projects = projects, selectedProject = projectId)
	else:
		return redirect(url_for('login'))

@app.route('/logout', methods=['GET'])
def logout():
    # remove the username from the session if it's there
    session.pop('username', None)
    return redirect(url_for('login'))
	
@app.route('/setup', methods=['GET'])
def setup():
	query = DB.get_case(active=request.args.get('active'), search=request.args.get('filter'))
	query+=DB.get_object(projectId=projectSession())
	return json.dumps(query)

@app.route('/design', methods=['POST'])
def design():
	return render_template('login.html')
	
#-----case page-----
@app.route('/case_page', methods=['GET'])
def case_page():
	query = DB.get_case(projectId = projectSession(),active=1, update=0)
	count = len(query)
	return render_template('case.html', cases=query, count=count)
	#return json.dumps(query)

@app.route('/caseForm/<int:projectId>', methods=['GET'])
def caseForm(projectId):
	query = DB.getAreasWithProject(projectId=projectId)
	return render_template('case.html', caseForm=query,count=len(query))
	
@app.route('/save_case', methods=['POST'])
def save_case():
	ID=DB.save_case(title=request.form["title"],priority=request.form["priority"],data=request.form["data"],area=request.form.getlist('areaBox'),projectId=projectSession(),dynamic=request.form.getlist('dynamicArea'))
	DB.save_steps(id=ID,action = request.form.getlist('action[]'),result = request.form.getlist('result[]'),projectId=projectSession())
	return json.dumps(ID)

@app.route('/load_case/<int:ID>/<mode>', methods=['GET'])
def load_case(ID,mode):
	query = DB.get_case_parameters(id=ID, projectId=projectSession())
	areaInCase = DB.getCaseArea(caseId=query[0][0])
	areas = DB.getAreas(projectId=projectSession())
	dynArea=DB.getDynamicAreas(projectId=projectSession())
	files = DB.getCaseFiles(caseId=ID)
	temp=[]
	boolen='false'
	for k in areas:
		for j in areaInCase:
			if j[0] == k[0]:
				boolen='true'
				casearea=j[2]
		if boolen=='false':
			temp.append([k,'notchecked',""])
		else:
			temp.append([k,'checked',casearea])
			boolen='false'
	if mode == "loadCase":
		return render_template('case.html', loadcase=query, areaInCase=areaInCase,areas=temp,count=len(areas),files=files)
	elif mode == "editCase":
		return render_template('case.html', editablecase=query, areaInCase=areaInCase,areas=temp,count=len(areas),files=files,dynArea=dynArea)	
	elif mode == "copy":
		return render_template('case.html', copycase=query, areaInCase=areaInCase,areas=temp,count=len(areas),files=files,dynArea=dynArea)	

@app.route('/get_step/<int:ID>/<mode>', methods=['GET'])
def get_step(ID,mode):
	query=DB.get_step_parameters(id=ID,projectId=projectSession())
	if query == False:
		return False
	else:
		stepIds=[]
		for k in query:
			stepIds=k[0]
		#pics=DB.getStepPics(stepIds=stepIds)
		#pics=list(pics)
		query=list(query)
		iterator=0
		if mode == "get_step":
			#for k in query:
			#	k=list(k)
			#	for l in pics:
			#		for j in l:
			#			if j[5] in k[1]:
			#				string = "<img src='"+j[2]+"' alt='"+j[3]+"' class='img-rounded' style='max-height:140px;max-width:140px;'>"
			#				k[1]=k[1].replace(j[5], string)
			#				j=['<-default->','<-default->','<-default->','<-default->','<-default->','<-default->']
			#			if j[5] in k[2]:
			#				string = "<img src='"+j[2]+"' alt='"+j[3]+"' class='img-rounded' style='max-height:140px;max-width:140px;'>"
			#				k[2]=k[2].replace(j[5], string)
			#				j=['<-default->','<-default->','<-default->','<-default->','<-default->','<-default->']
			#	query[iterator]=k
			#	iterator=iterator+1
			return render_template('step.html', step=query)
		elif mode == "editStep":
			return render_template('step.html', editablestep=query) #, pics=pics[0])
		
@app.route('/getStep/<int:caseId>', methods=['GET'])
def getStep(caseId):
	query=DB.get_step_parameters(id=caseId,projectId=projectSession())
	return json.dumps(query)
		
#---Fizikai törlés
@app.route('/deleteCasePhysical/<int:ID>', methods=['GET'])
def deleteCasePhysical(ID):
	DB.deleteCase(id=ID, projectId=projectSession())
	cases = DB.get_case(projectId = projectSession(),active=1,update=1)
	cases += DB.get_case(projectId = projectSession(),active=0,update=0)
	return render_template('admin.html', newCaseRequest="true", caseData=cases)
	
@app.route('/deleteAllCase', methods=['GET'])
def deleteAllCase():
	cases = DB.get_case(projectId = projectSession(),active=1,update=1)
	cases += DB.get_case(projectId = projectSession(),active=0,update=0)
	for k in cases:
		DB.deleteCase(id=k[0], projectId=projectSession())
	return render_template('admin.html', emptyCaseMan="true")
	
#---Logikai törlés
@app.route('/deleteCase/<int:ID>', methods=['GET'])
def deleteCase(ID):
	#DB.deleteCase(id=ID, projectId=projectSession())
	DB.deleteCaseLogic(id=ID, projectId=projectSession())
	return "ok"

#---Fizikai update	
@app.route('/updateCasePhysical/<int:ID>', methods=['POST'])
def updateCasePhysical(ID):
	DB.updateCase(title=request.form["title"],priority=request.form["priority"],data=request.form["data"],caseId=ID,area=request.form.getlist('areaBox'))
	DB.updateStep(caseId=ID,action = request.form.getlist('action[]'),result = request.form.getlist('result[]'))
	return str(ID)

#---Logikai update
@app.route('/updateCase/<int:ID>', methods=['POST'])
def updateCase(ID):
	DB.updateCaseLogic(caseId=ID)
	newCaseId=DB.save_case(title=request.form["title"],priority=request.form["priority"],data=request.form["data"],area=request.form.getlist('areaBox'),projectId=projectSession(),dynamic=request.form.getlist('dynamicArea'))
	DB.caseUpdateFlag(oldCaseId=ID,newCaseId=newCaseId)
	DB.save_steps(id=newCaseId,action = request.form.getlist('action[]'),result = request.form.getlist('result[]'),projectId=projectSession())
	return json.dumps(newCaseId)

#---Export
@app.route('/exportCaseToWord/<int:ID>', methods=['GET'])
def exportCaseToWord(ID):
	try:
		case=DB.get_case_parameters(id=ID,projectId=projectSession())
		steps=DB.get_case_step(caseId=ID,projectId=projectSession())
		name="./downloads/"+case[0][1].replace(".","_").replace("/","-")+".docx"
		DB.insertFile(name=name,projectId=projectSession())
	except:
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in exportCaseToWord method
			""")
		return False
	try:
		document = Document()
		document.add_heading(case[0][1], 0)
		p = document.add_paragraph(case[0][3])

		document.add_heading('Steps:', level=1)

		table = document.add_table(rows=1, cols=2)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Action'
		hdr_cells[1].text = 'Result'
		for item in steps:
			row_cells = table.add_row().cells
			row_cells[0].text = render_template("caseToDoc.html", caseDoc=item[3])
			row_cells[1].text = render_template("caseToDoc.html", caseDoc=item[4])
		document.save(name)
	except:
		print("""
			Unexpected error: Unexcepted error occured during the docx writing method in exportCaseToWord
			""")
		return False
	return name	

@app.route('/exportCaseToPDF/<int:ID>/<int:templateid>', methods=['GET'])
def exportCaseToPDF(ID,templateid):
	try:
		case=DB.get_case_parameters(id=ID,projectId=projectSession())
		steps=DB.get_case_step(caseId=ID,projectId=projectSession())
		template=UP.getTemplateFile(fileId=templateid)
		path='pdfTemplates/'+str(template[2])
		html=render_template(path,casePDF=case,steps=steps)
		filename="./downloads/"+case[0][1].replace(".","_").replace("/","-")+".pdf"
		DB.insertFile(name=filename,projectId=projectSession())
	except:
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in exportCaseToPDF method
			""")
		return False
	try:
		path=bytes(r'./wkhtmltopdf/bin/wkhtmltopdf.exe','utf-8')
		config = pdfkit.configuration(wkhtmltopdf=path)
		pdfkit.from_string(html, filename, configuration = config)
	except:
		print("""
			Unexpected error: Unexcepted error occured during the pdf writing method in exportCaseToPDF
			""")
		exc_type, exc_obj, exc_tb = sys.exc_info()
		fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
		f = exc_tb.tb_frame
		lineno = exc_tb.tb_lineno
		linecache.checkcache(fname)
		line = linecache.getline(fname, lineno, f.f_globals)
		print(exc_type, fname, lineno)
		print(line)
		return False
	return filename

@app.route('/exportCaseToXLSX/<int:ID>', methods=['GET'])
def exportCaseToXLSX(ID):
	try:
		case=DB.get_case_parameters(id=ID,projectId=projectSession())
		steps=DB.get_case_step(caseId=ID,projectId=projectSession())
		filename="./downloads/"+case[0][1].replace(".","_").replace("/","-")+".xlsx"
		DB.insertFile(name=filename,projectId=projectSession())
	except:
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in exportCaseToXLSX method
			""")
	try:
		workbook = xlsxwriter.Workbook(filename)
		worksheet = workbook.add_worksheet()
		bold = workbook.add_format({'bold': True})
		worksheet.set_column('A:D', len(steps)+4)
		worksheet.write('A1', 'Title:', bold)
		worksheet.write('A2', 'Description:', bold)
		worksheet.write('B1', case[0][1])
		worksheet.write('B2', case[0][3])
		worksheet.write('A4', 'Action', bold)
		worksheet.write('B4', 'Result', bold)
		for id,item in enumerate(steps):
			worksheet.write(id+4, 0, render_template("caseToDoc.html", caseDoc=item[3]))
			worksheet.write(id+4, 1, render_template("caseToDoc.html", caseDoc=item[4]))
		workbook.close()
	except:
		print("Unexpected error: Unexcepted error occured during the worksheet method in exportCaseToXLSX")
	return filename	
	
#-----object page-----	
@app.route('/object_page', methods=['GET'])
def object_page():
	query = DB.get_object(projectId=projectSession(),active=1)
	return render_template("object.html", object=query)

@app.route('/objectForm/<int:projectId>', methods=['GET'])
def objectForm(projectId):
	query = DB.getAreas(projectId=projectId)
	return render_template('object.html', objectForm=query,count=len(query))
	
@app.route('/load_object/<int:ID>/<mode>', methods=['GET'])
def load_object(ID,mode):
	query = DB.get_object_parameters(id=ID, projectId=projectSession())
	areaInObject = DB.getObjectArea(objectId=ID)
	areas = DB.getAreas(projectId=projectSession())
	files = DB.getObjectFiles(obId=ID)
	if files == []:
		files="empty"
	temp=[]
	boolen='false'
	for k in areas:
		for j in areaInObject:
			if j[0] == k[0]:
				boolen='true'
		if boolen=='false':
			temp.append([k,'notchecked'])
		else:
			temp.append([k,'checked'])
			boolen='false'
	if mode == "loadObject":
		return render_template("object.html", loadObject=query,areas=temp,count=len(areas),files=files)
	elif mode == "editObject":
		return render_template("object.html", editObject=query,areas=temp,count=len(areas),files=files)

@app.route('/save_object', methods=['POST'])	
def save_object():
	ID=DB.save_object(name=request.form["name"],hardware=request.form["hardware"],desc=request.form["desc"],version=request.form["version"],projectId=projectSession(),areas=request.form.getlist('areaBox'))
	return json.dumps(ID)

@app.route('/loadObjectModal/<int:id>', methods=['GET'])	
def loadObjectModal(id):
	files=DB.getObjectFiles(obId=id)
	return render_template("object.html", modalFiles=files)

@app.route('/updateObject', methods=['POST'])	
def updateObject():
	DB.updateObject(objectId=request.form["objectId"],name=request.form["name"],hardware=request.form["hardware"],desc=request.form["desc"],version=request.form["version"],projectId=request.form["projectId"],areas=request.form.getlist('areaBox'))
	return json.dumps(request.form["objectId"])
	
#--- Fizikai törlés	
@app.route('/deleteObjectPhysical/<int:ID>', methods=['GET'])
def deleteObjectPhysical(ID):
	DB.deleteObject(id=ID,projectId=projectSession())
	query = DB.get_object(projectId=projectSession(),active=0)
	return render_template('admin.html', newObjectRequest="true", objectData=query)

@app.route('/deleteAllObject', methods=['GET'])
def deleteAllObject():
	query = DB.get_object(projectId=projectSession(),active=0)
	for k in query:
		DB.deleteObject(id=k[0],projectId=projectSession())
	return render_template('admin.html', emptyObMan="true")
	
#--- Logikai törlés
@app.route('/deleteObject/<int:ID>', methods=['GET'])
def deleteObject(ID):
	DB.deleteObjectLogical(id=ID,projectId=projectSession())
	return "ok"
	
#-----set page-----	
@app.route('/set_page', methods=['GET'])
def set_page():
	query = DB.get_set(projectId=projectSession(),active=1,update=0)
	return render_template('set.html', set=query)

@app.route('/setForm', methods=['GET'])
def setForm():
	areas=DB.getAreas(projectId=projectSession())
	return render_template('set.html', setForm='true',areas=areas,count=len(areas))

@app.route('/save_set', methods=['POST'])	
def save_set():
	setId=DB.save_set(name=request.form["name"],date=request.form["date"],priority=request.form["priority"],projectId=projectSession(),areas=request.form.getlist('areaBox'))
	DB.saveSetCase(ID=request.form.getlist('ID'),setID=setId)
	return json.dumps(setId)

#--- Fizikai update
@app.route('/updateSetPhysical', methods=['POST'])	
def updateSetPhysical():
	DB.updateSet(setId=request.form['setId'],ID=request.form.getlist('ID'),name=request.form["name"],date=request.form["date"],priority=request.form["priority"],projectId=projectSession(),areas=request.form.getlist('areaBox'))
	return "OK"
	
#--- Logikai update
@app.route('/updateSet', methods=['POST'])	
def updateSet():
	DB.updateSetLogical(setId=request.form['setId'])
	setId=DB.save_set(name=request.form["name"],date=request.form["date"],priority=request.form["priority"],projectId=projectSession(),areas=request.form.getlist('areaBox'))
	DB.setUpdateFlag(oldSetId=request.form['setId'],newSetId=setId)
	DB.saveSetCase(ID=request.form.getlist('ID'),setID=setId)
	return json.dumps(setId)
	
@app.route('/load_set/<int:ID>/<mode>', methods=['GET'])
def load_set(ID,mode):
	query=DB.get_set_parameters(id=ID)
	cases=DB.getSetCases(id=ID)
	areaInSet = DB.getSetArea(setId=ID)
	areas = DB.getAreas(projectId=projectSession())
	files=DB.getSetFiles(setId=ID)
	temp=[]
	boolen='false'
	for k in areas:
		for j in areaInSet:
			if j[0] == k[0]:
				boolen='true'
		if boolen=='false':
			temp.append([k,'notchecked'])
		else:
			temp.append([k,'checked'])
			boolen='false'
	if mode == "loadSet":
		return render_template('set.html', loadSet=query, loadCase=cases,areas=temp,count=len(areas),files=files)
	if mode == "exeCasesBySet":
		return render_template('set.html', exeCasesBySet=cases)
	else:
		return render_template('set.html', loadEditableSet=query, editCase=cases,areas=temp,count=len(areas),files=files)

#--- Fizikai törlés
@app.route('/deleteSetPhysical/<int:ID>', methods=['GET'])
def deleteSetPhysical(ID):
	DB.deleteSet(id=ID)
	sets = DB.get_set(projectId=projectSession(),active=0,update=0)
	sets += DB.get_set(projectId=projectSession(),active=1,update=1)
	return render_template('admin.html',emptySetMan="true")

@app.route('/deleteAllSet', methods=['GET'])
def deleteAllSet():
	sets=DB.get_set(projectId=projectSession(),active=0,update=0)
	sets += DB.get_set(projectId=projectSession(),active=1,update=1)
	for k in ids:
		DB.deleteSet(k[0])
	return render_template('admin.html',newSetRequest="true",setData=sets)
	
#--- Logikai törlés	
@app.route('/deleteSet/<int:ID>', methods=['GET'])
def deleteSet(ID):
	DB.deleteSetLogical(id=ID)
	return "OK"

#---Export
@app.route('/exportSetToWord/<int:ID>', methods=['GET'])
def exportSetToWord(ID):
	set=DB.get_set_parameters(id=ID)
	cases=DB.getSetCases(id=set[0][0])
	document = Document()
	document.add_heading(set[0][1], 0)
	document.add_paragraph("")
	#document.add_page_break()
	for i,k in enumerate(cases):
		document.add_heading(k[1], level=1)
		description=document.add_paragraph("")
		description.add_run(k[3]).italic=True
		table = document.add_table(rows=1, cols=2)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Action'
		hdr_cells[1].text = 'Result'
		steps=DB.get_case_step(caseId=k[0],projectId=projectSession())
		for item in steps:
			row_cells = table.add_row().cells
			row_cells[0].text = render_template("caseToDoc.html", caseDoc=item[3])
			row_cells[1].text = render_template("caseToDoc.html", caseDoc=item[4])
		if i < len(cases)-1:
			document.add_page_break()
	filename="./downloads/"+set[0][1].replace(".","_").replace("/","-")+".docx"
	DB.insertFile(name=filename,projectId=projectSession())
	document.save(filename)
	return filename

@app.route('/exportSetToPDF/<int:ID>/<int:templateid>', methods=['GET'])
def exportSetToPDF(ID,templateid):
	try:
		set=DB.get_set_parameters(id=ID)
		cases=DB.getSetCases(id=set[0][0])
		template=UP.getTemplateFile(fileId=templateid)
		steps=[]
		for k in cases:
			steps.append(DB.get_case_step(caseId=k[0],projectId=projectSession()))
		path='pdfTemplates/'+str(template[2])
		html=render_template(path,setPDF=set,cases=cases,steps=steps)
		filename="./downloads/"+set[0][1].replace(".","_").replace("/","-")+".pdf"
		DB.insertFile(name=filename,projectId=projectSession())
	except Exception as e:
		exc_type, exc_obj, exc_tb = sys.exc_info()
		fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in exportSetToPDF method in
			""")
		print(exc_type, fname, exc_tb.tb_lineno)
		return False
	try:
		path=bytes(r'./wkhtmltopdf/bin/wkhtmltopdf.exe','utf-8')
		config = pdfkit.configuration(wkhtmltopdf=path)
		pdfkit.from_string(html, filename, configuration=config)
	except:
		print("""
			Unexpected error: Unexcepted error occured during the pdf writing method in exportSetToPDF
			""")
		return False
	return filename		
	
#-----execution-----	
@app.route('/execution_page', methods=['GET'])
def execution_page():
	query=DB.getExeOBTest(projectId=projectSession())
	return render_template('execution.html', execution=query)

@app.route('/newExe', methods=['GET'])
def newExecution():
	query=DB.get_object(projectId=projectSession(),active=1)
	areas=DB.getAreas(projectId=projectSession())
	users=DB.getUsers()
	return render_template('execution.html', newExe=query,areas=areas,count=len(areas),users=users,username=session['username'])	

@app.route('/saveExe', methods=['GET', 'POST'])
def SaveExecution():
	exeId=DB.saveExe(name=request.form["title"],testObject=request.form["TO"],projectId=projectSession(),areas=request.form.getlist('areaBox'),userId=request.form["userId"],dynamic=request.form.getlist('dynamicArea'))
	DB.saveCaseExe(ID=request.form.getlist('ID'),exeID=exeId)
	return json.dumps(exeId)

@app.route('/updateExe', methods=['POST'])
def UpdateExecution():
	DB.updateExecution(exeId=request.form['exeId'],name=request.form["title"],testObject=request.form["TO"],projectId=request.form["projectId"],areas=request.form.getlist('areaBox'),userId=request.form["userId"],dynamic=request.form.getlist('dynamicArea'))
	DB.updateCaseExe(ID=request.form.getlist('ID'),exeId=request.form['exeId'])
	return json.dumps(request.form['exeId'])
	
@app.route('/loadExecution/<int:ID>/<mode>', methods=['GET'])
def loadExecution(ID,mode):
	query=DB.getExeParameters(id=ID)
	object=DB.getExeObject(id=ID)
	users=DB.getUsers()
	cases=DB.getExeCases(id=ID)
	objects=DB.get_object(projectId=projectSession(), notRes=object[0],active=1)
	areaInExe = DB.getExeArea(exeId=ID)
	areas = DB.getAreasWithProject(projectId=projectSession())
	dynArea=DB.getDynamicAreas(projectId=projectSession())
	files = DB.getExeFiles(exeId=ID)
	temp=[]
	boolen='false'
	for k in areas:
		for j in areaInExe:
			if j[0] == k[0]:
				boolen='true'
				exearea=j[2]
		if boolen=='false':
			temp.append([k,'notchecked',""])
		else:
			temp.append([k,'checked',exearea])
			boolen='false'
	if mode == "loadExe":
		return render_template('execution.html', loadExe=query, loadCase=cases, loadObject=object,areas=temp,count=len(areas),exeId=ID,files=files,users=users)
	if mode == "editExe":
		return render_template('execution.html', loadEditableExe=query, loadEditableCase=cases, loadEditableObject=object, loadObjects=objects,areas=temp,count=len(areas),exeId=ID,files=files,users=users,dynArea=dynArea)
	#else:
	#	return render_template('execution.html', loadEditableSet=query, editCase=cases)

@app.route('/deleteExe/<int:ID>/<int:obID>', methods=['GET'])
def deleteExe(ID,obID):
	DB.deleteExe(id=ID,obid=obID)
	DB.deleteFilesInTest(exeId=ID)
	return "ok"

@app.route('/getFirstCaseID/<int:id>', methods=['GET'])
def getFirstCaseID(id):
	query = DB.getExeFromCases(id=id)
	return json.dumps(query[0][2])
	
#-----Report-----	

@app.route('/getReports', methods=['GET'])
def getReports():
	query = Report.getRecords()
	return render_template('dashboard.html', dashboard=query, user=session['username'])

@app.route('/save_report', methods=['POST'])	
def save_report():
	Report.save_report(title=request.form["title"],user=request.form["user"],date=request.form["date"],report=request.form["desc"])
	query = Report.getRecords()
	return render_template('dashboard.html', dashboard=query, user=session['username'])
	
@app.route('/getReporttoEdit/<int:id>', methods=['GET'])
def getReporttoEdit(id):
	query = Report.getRepPar(ID=id)
	return render_template("dashboard.html", editableReport=query)

@app.route('/getReporttoLoad', methods=['POST'])
def getReporttoLoad():
	query = Report.getRecords()
	return render_template("dashboard.html", dashboard=query, user=session['username'])

@app.route('/save_edited_record/<int:id>', methods=['POST'])	
def save_edited_record(id):
	Report.save_edited_record(ID=id,title=request.form["title"],user=request.form["user"],date=request.form["date"],report=request.form["desc"])
	query = Report.getRecords()
	return render_template('dashboard.html', dashboard=query, user=session['username'])
	
	
@app.route('/loadSearchForm', methods=['GET'])
def loadSearchForm():
	query = DB.getUsers();
	return render_template('search.html', users=query)

#-----Test-----

@app.route('/testSetup', methods=['GET'])
def testSetup():
	object=DB.getExeOBTest(projectId=projectSession())
	return render_template('test.html', Exe=object)
	
@app.route('/loadTest/<int:id>', methods=['GET'])
def loadTest(id):
	query = DB.getExeFromCases(id=id)
	return render_template('test.html', loadTest=query)

@app.route('/testPage/<int:id>/<int:exeId>', methods=['GET'])
def testPage(id,exeId):
	query = DB.get_step_parameters(id=id,projectId=projectSession())
	res = DB.getStatusFromStepExe(exeId=exeId, caseId=id)
	query=zip(query,res)
	return render_template('test.html', step=query, status=res,exeId=exeId)

@app.route('/saveStatus/<int:stepId>/<int:caseExeId>/<status>', methods=['GET'])
def saveStatus(stepId,caseExeId,status):
	DB.saveSetStatus(stepId=stepId,caseExeId=caseExeId,status=status)
	return "Ok"

@app.route('/saveCaseStatus/<int:exeId>/<int:caseId>', methods=['GET'])
def saveCaseStatus(exeId,caseId):
	status=DB.saveCaseStatus(exeId=exeId,caseId=caseId)
	return render_template("test.html", status=status)

@app.route('/saveComment/<int:stepId>/<int:exeId>', methods=['POST'])
def saveComment(stepId,exeId):
	if request.form['comment']=="":
		comment="NULL"
	else:
		comment=request.form['comment']
	DB.saveComment(stepId=stepId,exeId=exeId,comment=comment)
	return comment

@app.route('/updateStepExeCorrectWay', methods=['GET'])
def updateStepExeCorrectWay():
	DB.updateStepExeCorrectWay()
	return "ok"
	
@app.route('/loadModal/<int:stepId>/<int:exeId>', methods=['GET'])
def loadModal(stepId,exeId):
	#step=DB.getStepExeParam(stepId=stepId,exeId=exeId)
	var=DB.getVariables(projectId=projectSession())
	stepVar=DB.getExeStepVar(stepId=stepId,exeId=exeId)
	temp=[]
	steps=""
	for k in var:
		if stepVar == None:
			temp.append([k,None])
		else:
			boolen=False
			for l in stepVar:
				if int(l[1])==int(k[0]):
					boolen=True
					steps=l
			if boolen==False:
				temp.append([k,None])
			else:
				temp.append([k,steps])
	modalSteps=[stepId,exeId]
	return render_template("test.html",modalSteps=modalSteps,variables=temp)

@app.route('/sendVarTest/<int:exeId>/<int:stepId>/<int:varId>', methods=['POST'])
def sendVarTest(exeId,stepId,varId):
	DB.saveVarTest(stepId=stepId,exeId=exeId,varId=varId,value=request.form['value'])
	return "OK"

@app.route('/clearValue/<int:exeId>/<int:stepId>/<int:varId>', methods=['GET'])
def clearValue(exeId,stepId,varId):
	DB.clearValue(stepId=stepId,exeId=exeId,varId=varId)
	return "OK"
	
#-----Project-----
@app.route('/projectChanging/<int:id>', methods=['GET'])
def projectChanging(id):
	DB.setProjectToUser(id=id, user=session['username'])
	return "OK"	

def projectSession():
	return DB.getSelectedProject(user=session['username'])
	
#-----Chart-----
@app.route('/requestChart/<type>/<int:projectId>/<int:limit>', methods=['GET'])
def requestChart(type,projectId,limit):
	if type=="line":
		return chartReload(type,"All",0,0,"All",limit,0,projectId)
	else:
		query = DB.getDataForCharts(projectId=projectId,interval=limit)
		passed=0
		failed=0
		skipped=0
		notimp=0
		all=0
		for k in query:
			if k[2] == "RUN":
				passed+=1
			if k[2] == "FAILED":
				failed+=1
			if (k[2] == "NOTRUN") or (k[2] == "SKIPPED"):
				skipped+=1
			if k[2] == "NOTIMP":
				notimp+=1
			all+=1
		rate=[passed,failed,skipped,notimp,all]
		return render_template('test2.html',type=type,rate=rate).replace('\n','')
	
@app.route('/chartFilter/<type>/<int:projectId>', methods=['GET'])
def chartFilter(type,projectId):
	query = DB.getChartFilterData(projectId=projectId)
	executions = DB.getExeOBTest(projectId=projectSession())
	areas = DB.getAreas(projectId=projectId)
	return render_template('chartFilter.html',type=type,data=query,areas=areas,executions=executions)

@app.route('/chartExeReload/<int:obId>', methods=['GET'])
def chartExeReload(obId):
	query = DB.reloadChartFilterData(obId=obId,projectId=projectSession())
	return render_template('chartFilter.html',reloadExe=query)
	
@app.route('/chartReload/<type>/<interval>/<int:obId>/<int:areaId>/<status>/<int:limit>/<int:exeId>/<int:projectId>', methods=['GET'])
def chartReload(type,interval,obId,areaId,status,limit,exeId,projectId):
	passed=0
	failed=0
	skipped=0
	notimp=0
	all=0
	result = DB.getFilteredPar(objectId=obId,areaId=areaId,status=status,interval=limit,exeId=exeId,projectId=projectId)
	render=""
	if result:
		if type == "pie":
			for k in result:
				if k[0] == "RUN":
					passed+=1
				if k[0] == "FAILED":
					failed+=1
				if (k[0] == "NOTRUN") or (k[0] == "SKIPPED"):
					skipped+=1
				if k[0] == "NOTIMP":
					notimp+=1
				all+=1
				rate=[passed,failed,skipped,notimp,all]
				render = render_template('test2.html',type=type,rate=rate)
		if type == "line":
			temp=[]
			container=[]
			boolen = "false"
			for k in result:
				if k[0] != 'default':
					temp.append(k)
				default = k
				for j in result:
					if (default[3] == j[3]) and (default[4] != j[4]):
						temp.append(j)
						index=result.index(j)
						result.pop(index)
						result.insert(index,('default','default','default','default','default'))
						boolen="true"
				if temp:
					passed=0
					skipped=0
					failed=0
					notimp=0
					all=0
					for o in temp:
						if o[0] != "default":
							if o[0] == "RUN":
								passed+=1
							if o[0] == "FAILED":
								failed+=1
							if (o[0] == "NOTRUN") or (o[0] == "SKIPPED"):
								skipped+=1
							if o[0] == "NOTIMP":
								notimp+=1
							all+=1
				rate=[passed,failed,skipped,notimp,all]
				if boolen == "false":
					index=result.index(k)
					result.pop(index)
					result.insert(index,('default','default','default','default','default'))
				else:
					boolen = "false"
				if temp:
					if (temp[0][0] != "default"):
						container.append([temp,rate])
				result.pop(0)
				result.insert(0,('default','default','default','default','default'))
				temp = []
			render = render_template('test2.html',type=type,data=container)
		return render
	else:
		if type=="pie":
			return render_template('test2.html',result="pie")
		else:
			return render_template('test2.html',result="line")
	
@app.route('/jenkinsRadiator/<int:limit>', methods=['GET'])
def jenkinsRadiator(limit):
	query = DB.getJenkinsData(projectId=projectSession(),limit=limit)
	cases = DB.getJenkinsCaseResult(data=query,projectId=projectSession())
	rendered = render_template('jenkinsRadiator.html')
	if len(cases) > 0:
		default = [cases[0]]
		temp = []
		sorting = []
		for k in cases:
			if k[0] == default[0][0]:
				if k != default[0]:
					default.append(k)
			else:
				temp.append(default)
				default = [k]
		temp.append(default)
		default = [k]
		all=0
		passed=0
		skipped=0
		failed=0
		iterator=0
		for k in temp:
			for j in k:
				if j[1] == 'RUN':
					passed+=1
				if j[1] == 'SKIPPED':
					skipped+=1
				if j[1] == 'NOTRUN':
					skipped+=1
				if j[1] == 'FAILED':
					failed+=1
				all+=1
			rate = []
			rate.append(passed)
			rate.append(skipped)
			rate.append(failed)
			rate.append(all)
			for l in temp[iterator]:
				tupList=list(l)
				tupList[3]=tupList[3].encode('ascii', 'backslashreplace').decode("utf-8", "replace")
				l=tuple(tupList)
			tupList=list(query[iterator])
			tupList[1]=tupList[1].encode('ascii', 'backslashreplace').decode("utf-8", "replace")
			tupList[2]=tupList[2].encode('ascii', 'backslashreplace').decode("utf-8", "replace")
			tupList[4]=tupList[4].encode('ascii', 'backslashreplace').decode("utf-8", "replace")
			query[iterator]=tuple(tupList)
			if passed/all >= 0.8:
				#rendered += render_template('jenkinsRadiator.html', mode='success', param=temp[iterator], data=query[iterator], iterator=iterator, rate=rate)
				sorting.append(('success',temp[iterator],query[iterator],iterator,rate))
			else:
				if failed/all > 0:
					#rendered += render_template('jenkinsRadiator.html', mode='danger', param=temp[iterator], data=query[iterator], iterator=iterator, rate=rate)
					sorting.append(('danger',temp[iterator],query[iterator],iterator,rate))
				else:
					#rendered += render_template('jenkinsRadiator.html', mode='warning', param=temp[iterator], data=query[iterator], iterator=iterator, rate=rate)
					sorting.append(('warning',temp[iterator],query[iterator],iterator,rate))
			all=0
			passed=0
			iterator+=1
			skipped=0
			failed=0
		temp = []
		iterator=0
		b=0
		default1=[]
		boolen = "false"
		for k in sorting:
			temp = []
			temp.append(k)
			default=k[2][0]
			for l in sorting:
				if l[2][0] == default:
					if l[2] != k[2]:
						temp.append(l)
						index=sorting.index(l)
						sorting.pop(index)
						sorting.insert(index,('default','default',('default','default')))
						#default1=l
						boolen="true"
			passed=0
			skipped=0
			failed=0
			all=0
			for o in temp:
				if o[0] != "default":
					passed+=o[4][0]
					skipped+=o[4][1]
					failed+=o[4][2]
					all+=o[4][3]
			rate=[]
			rate.append(passed)
			rate.append(skipped)
			rate.append(failed)
			rate.append(all)
			if boolen == "false":
				index=sorting.index(k)
				sorting.pop(index)
				sorting.insert(index,('default','default',('default','default')))
			else:
				boolen = "false"
			if temp[0][0] != "default":
				rendered+=render_template('jenkinsRadiator.html', objectExe=temp, iterator=iterator, rate=rate)
			sorting.pop(0)
			sorting.insert(0,('default','default',('default','default')))
			iterator+=1
		if len(sorting) == 1 and sorting[0][0] != "default":
			passed=sorting[0][4][0]
			skipped=sorting[0][4][1]
			failed=sorting[0][4][2]
			all=sorting[0][4][3]
			rate=[]
			rate.append(passed)
			rate.append(skipped)
			rate.append(failed)
			rate.append(all)
			rendered+=render_template('jenkinsRadiator.html', objectExe=sorting, iterator=iterator, rate=rate)
	return rendered #render_template('jenkinsRadiator.html', jenkins=rendered)	

	
@app.route('/jenkinsFilter', methods=['GET'])
def jenkinsFilter():
	return render_template("jenkinsFilter.html")
	
#-----Dashboard-----
@app.route('/dashboardLoad', methods=['GET'])
def dashboardLoad():
	return render_template('slide.html')

#-----Admin----
@app.route('/getUser', methods=['GET'])
def getUser(active=None):
	return json.dumps(session['username'])

@app.route('/getUsers', methods=['GET'])
def getUsers():
	query = DB.getUsers();
	roles = DB.getRoles();
	return render_template('admin.html', admin=query,roles=roles)

@app.route('/getAdminNav', methods=['GET'])
def getAdminNav():
	return render_template('admin.html', adminNav="true")
	
@app.route('/savePassword', methods=['POST'])
def savePassword():
	return DB.updatePw(oldPw=request.form["oldPw"],newPw=request.form["newPw"],id=request.form['id'])

@app.route('/updateUserRole', methods=['POST'])
def updateUserRole():
	return DB.updateUserRole(userId=request.form['userId'],roleId=request.form['roleId'])
	
@app.route('/userActive', methods=['POST'])
def userActive():
	DB.userActive(userId=request.form['userId'],userStatus=request.form['status'])
	return render_template('admin.html', userStatus=request.form['status'])

@app.route('/saveUser', methods=['POST'])	
def saveUser():
	return DB.saveUser(userName=request.form['userName'],pw=request.form['password'],roleId=request.form['roleId'],projectId=request.form['projectId'])

@app.route('/deleteUser', methods=['POST'])	
def deleteUser():
	DB.deleteUser(userId=request.form['userId'])
	return "OK"

@app.route('/getProjectManagement', methods=['GET'])	
def getProjectManagement():
	projects = DB.getProjects()
	tags = DB.getAreasWithProject(projectId = projectSession())
	vars = DB.getVariables(projectId = projectSession())
	projects=DB.getProjects()
	return render_template('admin.html', projectManagement=projects, tags=tags, vars=vars, projects=projects)
	
@app.route('/projectActive', methods=['POST'])
def projectActive():
	DB.projectActive(projectId=request.form['projectId'],projectStatus=request.form['status'])
	return render_template('admin.html', projectStatus=request.form['status'])
	
@app.route('/deleteProject', methods=['POST'])
def deleteProject():
	DB.deleteProject(projectId=request.form['projectId'])
	DB.DeleteEmailConfig(projectId=request.form['projectId'])
	return "OK"
	
@app.route('/saveProject', methods=['POST'])	
def saveProject():
	id = DB.saveProject(projectName=request.form['projectName'])
	DB.CreateEmailConfig(projectId=id)
	return str(id)

@app.route('/saveVariable', methods=['POST'])	
def saveVariable():
	newvar=DB.saveVariable(variable=request.form['variable'],projectId=request.form['projectId'])
	if newvar == "false":
		return "false"
	else:
		vars = DB.getVariables(projectId = projectSession())
		projects=DB.getProjects()
		return render_template("admin.html",newvar=vars,projects=projects)

@app.route('/deleteVar', methods=['POST'])	
def deleteVar():
	DB.deleteVariable(variableId=request.form['varId'])
	return "OK"
		
@app.route('/getDatabaseManagement', methods=['GET'])	
def getDatabaseManagement():
	cases = DB.get_case(projectId = projectSession(),active=1,update=1)
	cases += DB.get_case(projectId = projectSession(),active=0,update=0)
	sets = DB.get_set(projectId=projectSession(),active=0,update=0)
	sets += DB.get_set(projectId=projectSession(),active=1,update=1)
	executions = DB.getExecution(projectId=projectSession())
	objects = DB.get_object(projectId=projectSession(),active=0)
	temp = [cases,sets,executions,objects]
	return render_template('admin.html', databaseManagement=temp)
	
@app.route('/getExportImport', methods=['GET'])	
def getExportImport():
	cases = DB.get_case(projectId = projectSession(),active=1,update=0)
	sets = DB.get_set(projectId=projectSession(),active=1,update=0)
	executions = DB.getExeOBTest(projectId=projectSession())
	objects = DB.get_object(projectId=projectSession(),active=1)
	templates=UP.getTemplates()
	dbtables=DB.getTablesFromDB()
	files=DB.getDownloadFiles(projectId=projectSession())
	emails = DB.getEmails(projectId=projectSession())
	users= DB.getUsers(projectId=projectSession())
	email_config = DB.GetEmailConfig(projectId=projectSession())
	return render_template('exportimport.html', exportimport=True,cases=cases,sets=sets,exes=executions, \
		objects=objects,tables=dbtables,files=files,templates=templates,emails=emails, \
		users=users,email_config=email_config)

@app.route('/deleteTag', methods=['POST'])	
def deleteTag():
	DB.deleteArea(id=request.form['tagId'])
	return "OK"

@app.route('/saveTag', methods=['POST'])	
def saveTag():
	result=DB.saveTag(tagName=request.form['tagName'],dynamic=request.form['dynamic'],projectId=request.form['projectId'])
	tags=DB.getAreasWithProject(projectId=request.form['projectId'])
	if result == "false":
		return result
	else:
		return render_template("admin.html", newTag=tags)
	
@app.route('/downloadFiles', methods=['GET'])	
def downloadFiles():
	files=DB.getDownloadFiles(projectId=projectSession())
	return render_template("exportimport.html",downloadableFiles=files)

#----History----
@app.route('/HistoryForm', methods=['GET'])	
def HistoryForm():
	historyForm=DB.getExeOBTest(projectId=projectSession())
	return render_template('history.html', historyForm=historyForm)
	
@app.route('/historyExe/<mode>', methods=['GET'])	
def historyExe(mode):
	historyNav=DB.getExeOBTest(projectId=projectSession())
	return render_template('history.html', historyNav=historyNav,mode=mode)
	
@app.route('/loadHistoryExe/<int:exeId>/<exeStatus>', methods=['GET'])	
def loadHistoryExe(exeId,exeStatus):
	loadHistory=DB.getExeResult(exeId=exeId)
	comments=DB.getExeComments(exeId=exeId)
	files=DB.getExeStepFiles(exeId=exeId)
	exe=DB.getExeOBHist(projectId=projectSession(),exeId=exeId)
	return render_template('history.html', loadHistory=loadHistory, exeName=exe, status=exeStatus,files=files,comments=comments)
	
@app.route('/loadResultForm', methods=['GET'])	
def loaResultForm():
	sets=DB.get_set(projectId=projectSession(),active=1,update=0)
	return render_template('resultHistory.html', resultForm="true",sets=sets)
	
@app.route('/loadLastResultHist/<int:limit>/<int:setId>', methods=['GET'])	
def loadLastResultHist(limit,setId):
	caseIds=DB.getSetCases(id=setId)
	exes=DB.getLastExes(limit=limit,caseIds=caseIds)
	titles=DB.getExeTitles(exeIds=exes)
	result=DB.getCaseHistory(exeIds=exes)
	return render_template('resultHistory.html', loadLastResultHist=result, exes=exes, title=titles)
	
@app.route('/loadResStepForm', methods=['GET'])	
def loadResStepForm():
	sets=DB.get_set(projectId=projectSession(),active=1,update=0)
	return render_template('resultHistory.html', resultStepForm="true",sets=sets)
	
@app.route('/loadLastResStepHist/<int:limit>/<int:setId>', methods=['GET'])	
def loadLastResStepHist(limit,setId):
	caseIds=DB.getSetCases(id=setId)
	exes=DB.getLastExes(limit=limit,caseIds=caseIds)
	titles=DB.getStepTitles(exeIds=exes)
	result=DB.getStepHistory(exeIds=exes)
	return render_template('resultHistory.html', loadLastResStepHist=result, exes=exes, title=titles)
	
@app.route('/CaseHistForm', methods=['GET'])	
def CaseHistForm():
	sets=DB.get_set(projectId=projectSession(),active=1,update=0)
	return render_template('caseHistory.html', caseHistForm=sets)
	
@app.route('/historyCase/<int:setId>', methods=['GET'])	
def historyCase(setId):
	caseIds=DB.getSetCases(id=setId)
	exes=DB.getExeOBLimit(projectId=projectSession(),caseIds=caseIds,limit=5)
	result=DB.getResultCases(exes=exes)
	it1=0
	cases=[]
	for k in caseIds:
		res=DB.get_case_name(caseId=k[0])
		cases.append([k[0],res[0]])
		it1=it1+1
	return render_template('caseHistory.html', resultDiagrams=result,exes=exes,cases=caseIds)

@app.route('/loadCaseHistory/<int:caseId>', methods=['GET'])	
def loadCaseHistory(caseId):
	cases=DB.getCaseResHist(caseId=caseId)
	exes=DB.getExesForCases(cases=cases)
	return render_template('caseHistory.html', caseStatus=cases,exes=exes)

@app.route('/generatorForm', methods=['GET'])	
def generatorForm():
	tags=DB.getAreas(projectId=projectSession())
	vars=DB.getVariables(projectId=projectSession())
	return render_template("reportGenerator.html", generatorForm="true",tags=tags,vars=vars)	

@app.route('/getExeResultToGen/<int:id>', methods=['GET'])	
def getExeResultToGen(id):
	#all=DB.getExeResultOb(exeId=id)
	result=DB.getExeOb(exeId=id)
	return render_template("reportGenerator.html", newSelectedAdvanced=result)
	
#----File Upload----
PIC_EXTENSIONS = set(['png', 'jpg', 'jpeg', 'gif','PNG','JPG','JPEG','GIF'])
DOC_EXTENSIONS = set(['doc', 'docx', 'csv', 'xlsx', 'xlt', 'xls'])
	
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS

@app.route('/upload_file_test/<int:Id>/<mode>', methods=['GET', 'POST'])
def upload_file_test(Id,mode):
	if mode == "test":
		UPLOAD_FOLDER = './uploads/test/'
	if mode == "object":
		UPLOAD_FOLDER = './uploads/object/'
	if mode == "set":
		UPLOAD_FOLDER = './uploads/set/'
	if mode == "exe":
		UPLOAD_FOLDER = './uploads/execution/'
	if mode == "case":
		UPLOAD_FOLDER = './uploads/case/'
	if mode == "templates":
		UPLOAD_FOLDER = './templates/pdfTemplates/'
	app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
	if request.method == 'POST':
		if request.form['name'] == '':
			return "No selected file"
		if allowed_file(request.form['name']):
			name=UP.nameExists(path=os.path.join(app.config['UPLOAD_FOLDER'], request.form['name']),mode=mode,name=request.form['name'],folder=app.config['UPLOAD_FOLDER'])
			if request.form['name'].rsplit('.', 1)[1] in PIC_EXTENSIONS:
				file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
				file.write(base64.b64decode(request.form['context'].split(',')[1]))
			elif request.form['name'].rsplit('.', 1)[1] in DOC_EXTENSIONS:
				file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
				file.write(base64.b64decode(request.form['context'].split(',')[1]))
			else:
				file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'w')
				file.write(request.form['context'].encode('ascii', 'backslashreplace').decode("utf-8", "replace"))
			file.close()
			if mode == "test":
				result=UP.saveTestFile(stepexeId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'],name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
				return render_template("upload.html", addPlusToTest=result)
			if mode == "object":
				result=UP.saveObjectFile(objectId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
				return render_template("upload.html", objectFile=result)
			if mode == "set":
				result=UP.saveSetFile(setId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
				return render_template("upload.html", setFile=result)
			if mode == "exe":
				result=UP.saveExeFile(exeId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
				return render_template("upload.html", exeFile=result)
			if mode == "case":
				result=UP.saveCaseFile(caseId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
				return render_template("upload.html", caseFile=result)
			if mode == "templates":
				result=UP.saveTemplatesFile(url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1],status=1)
				temps=UP.getTemplates()
				return render_template("exportimport.html", templatesFile=temps)
	return "error"

@app.route('/upload_step_files/<int:Id>/<mode>/<replaceTag>', methods=['GET', 'POST'])
def upload_step_files(Id,mode,replaceTag):
	if mode == "action":
		UPLOAD_FOLDER = './uploads/step/action/'
	if mode == "result":
		UPLOAD_FOLDER = './uploads/step/result/'
	app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
	if request.method == 'POST':
		if request.form['name'] == '':
			return "No selected file"
		if allowed_file(request.form['name']):
			name=UP.nameExists(path=os.path.join(app.config['UPLOAD_FOLDER'], request.form['name']),mode=mode,name=request.form['name'],folder=app.config['UPLOAD_FOLDER'])
			if request.form['name'].rsplit('.', 1)[1] in PIC_EXTENSIONS:
				file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
				file.write(base64.b64decode(request.form['context'].split(',')[1]))
			elif request.form['name'].rsplit('.', 1)[1] in DOC_EXTENSIONS:
				file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
				file.write(base64.b64decode(request.form['context'].split(',')[1]))
			else:
				file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'w')
				file.write(request.form['context'].encode('ascii', 'backslashreplace').decode("utf-8", "replace"))
			file.close()
			if mode == "action" or mode == "result":
				result=UP.saveStepFile(stepId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1],replaceTag=replaceTag,status=mode)
				return "OK"

@app.route('/upload_file_area', methods=['POST'])
def upload_file_area():
	UPLOAD_FOLDER = './uploads/step/'
	app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
	if request.method == 'POST':
		f=request.files.getlist("image")[0]
		if f.filename == '':
			return "No selected file"
		if allowed_file(f.filename):
			name=UP.nameExists(path=os.path.join(app.config['UPLOAD_FOLDER'], f.filename),mode="step",name=f.filename,folder=app.config['UPLOAD_FOLDER'])
			f.save(os.path.join(app.config['UPLOAD_FOLDER'], name))
			result=UP.saveStepFile(url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=f.filename.rsplit('.', 1)[1],status="step")
			link=result[1].split('/')[1]+"/"+result[1].split('/')[2]+"/"+result[1].split('/')[3]
			data="{'type': 'image\\/"+result[3].rsplit('.')[0]+"','name:'"+name+"','link': '"+link+"'"
			print(data)
			return jsonify(data=data)
		return None
	return None
				
@app.route('/fileProperty/<int:exeId>/<int:stepId>', methods=['GET'])	
def fileProperty(exeId,stepId):
	files=UP.getTestFiles(exeId=exeId,stepId=stepId)
	stepexeId=DB.getStepExeId(exeId=exeId,stepId=stepId)
	if files == []:
		files="empty"
	return render_template('upload.html', testUpload=files, stepId=stepId,stepexeId=stepexeId)

@app.route('/upload_file_update/<int:id>/<mode>', methods=['POST'])	
def upload_file_update(id,mode):
	if mode == "object":
		UPLOAD_FOLDER = './uploads/object/'
	if mode == "set":
		UPLOAD_FOLDER = './uploads/set/'
	if mode == "case":
		UPLOAD_FOLDER = './uploads/case/'
	app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
	if request.method == 'POST':
		if request.form['name'] == '':
			return "No selected file"
		if allowed_file(request.form['name']):
			name=UP.nameExists(path=os.path.join(app.config['UPLOAD_FOLDER'], request.form['name']),mode=mode,name=request.form['name'],folder=app.config['UPLOAD_FOLDER'])
			if mode == "object":
				result=DB.checkFileInObjects(objectId=id,filename=request.form['name'])
				if result == None:
					if request.form['name'].rsplit('.', 1)[1] in PIC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					elif request.form['name'].rsplit('.', 1)[1] in DOC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					else:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'w')
						file.write(request.form['context'].encode('ascii', 'backslashreplace').decode("utf-8", "replace"))
					file.close()
					result=UP.saveObjectFile(objectId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
					return render_template('upload.html', objectFile=result)
			if mode == "set":
				result=DB.checkFileInSets(setId=id,filename=request.form['name'])
				if result == None:
					if request.form['name'].rsplit('.', 1)[1] in PIC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					elif request.form['name'].rsplit('.', 1)[1] in DOC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					else:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'w')
						file.write(request.form['context'].encode('ascii', 'backslashreplace').decode("utf-8", "replace"))
					file.close()
					result=UP.saveSetFile(objectId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
					return render_template('upload.html', objectFile=result)
			if mode == "exe":
				result=DB.checkFileInExes(exeId=id,filename=request.form['name'])
				if result == None:
					if request.form['name'].rsplit('.', 1)[1] in PIC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					elif request.form['name'].rsplit('.', 1)[1] in DOC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					else:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'w')
						file.write(request.form['context'].encode('ascii', 'backslashreplace').decode("utf-8", "replace"))
					file.close()
					result=UP.saveExeFile(exeId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
					return render_template('upload.html', exeFile=result)
			if mode == "case":
				result=DB.checkFileInCases(caseId=id,filename=request.form['name'])
				if result == None:
					if request.form['name'].rsplit('.', 1)[1] in PIC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					elif request.form['name'].rsplit('.', 1)[1] in DOC_EXTENSIONS:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'wb')
						file.write(base64.b64decode(request.form['context'].split(',')[1]))
					else:
						file = open(os.path.join(app.config['UPLOAD_FOLDER'], name),'w')
						file.write(request.form['context'].encode('ascii', 'backslashreplace').decode("utf-8", "replace"))
					file.close()
					result=UP.saveCaseFile(caseId=Id,url=os.path.join(app.config['UPLOAD_FOLDER'], name),filename=name,extension=request.form['name'].rsplit('.', 1)[1])
					return render_template('upload.html', caseFile=result)
		return "ok"
	return "error"
	
@app.route('/deleteFiles/<int:fileId>/<mode>', methods=['GET'])	
def deleteFiles(fileId,mode):
	if mode == "test":
		file=UP.getTestFileURL(fileId=fileId)
		os.remove(file)
		return UP.deleteFileTest(fileId=fileId)
	if mode == "object":
		file=UP.getObjectFileURL(fileId=fileId)
		os.remove(file)
		return UP.deleteFileObject(fileId=fileId)
	if mode == "set":
		file=UP.getSetFileURL(fileId=fileId)
		os.remove(file)
		return UP.deleteFileSet(fileId=fileId)
	if mode == "exe":
		file=UP.getExeFileURL(fileId=fileId)
		os.remove(file)
		return UP.deleteFileExe(fileId=fileId)
	if mode == "case":
		file=UP.getCaseFileURL(fileId=fileId)
		os.remove(file)
		return UP.deleteFileCase(fileId=fileId)
	if mode == "action" or mode == "result":
		file=UP.getStepFileURL(fileId=fileId)
		os.remove(file)
		return UP.deleteFileStep(fileId=fileId)
	if mode == "templates":
		file=UP.getTemplateFileURL(fileId=fileId)
		os.remove(file)
		return UP.deleteFileTemplate(fileId=fileId)
	
@app.route('/deleteStepFile/<url>/<mode>', methods=['GET'])	
def deleteStepFile(url,mode):
	if mode == "step":
		os.remove(url)
		return UP.deleteFileStepUrl(url=url)

#---- Export/Import ----
@app.route('/loadExportForm/<mode>/<int:id>', methods=['GET'])	
def loadExportForm(mode,id):
	templates=UP.getTemplates()
	if mode == "set":
		set=DB.get_set_parameters(id=id)
		return render_template("exportimport.html", setid=id,setname=set[0][1],templates=templates)
	if mode == "case":
		case=DB.get_case_parameters(id=id,projectId=projectSession())
		return render_template("exportimport.html", caseid=id, casename=case[0][1],templates=templates)
	if mode == "exe":
		exe=DB.getExeParameters(id=id)
		ob=DB.getExeObject(id=id)
		return render_template("exportimport.html", exeid=id, exename=exe[0][1], objectname=ob[1],templates=templates)

@app.route('/exportResultToWord/<int:id>', methods=['GET'])
def exportResultToWord(id):
	try:
		exeResult=DB.getExeResult(exeId=id)
		exeOb=DB.getExeOBHist(projectId=projectSession(),exeId=id)
		name=exeOb[1] + " (" + exeOb[2] + ")"
		filename="./downloads/"+name.replace(".","_").replace("/","-")+".docx"
		DB.insertFile(name=filename,projectId=projectSession())
	except:
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in exportResultToWord method
			""")
		return False
	try:
		document = Document()
		header = exeOb[1] + " (" + exeOb[2] + ") Report"
		document.add_heading(header, 0)
		table = document.add_table(rows=1, cols=4)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Case name'
		hdr_cells[1].text = 'Result'
		hdr_cells[2].text = 'Assigned to'
		hdr_cells[3].text = 'Duration'
		for item in exeResult:
			row_cells = table.add_row().cells
			row_cells[0].text = item[1]
			row_cells[1].text = item[2]
			row_cells[2].text = ""
			row_cells[3].text = ""
		document.save(filename)
	except:
		print("""
			Unexpected error: Unexcepted error occured during the docx writing method in exportResultToWord
			""")
		return False
	return filename
	
@app.route('/exportResultToPDF/<int:id>/<int:templateid>', methods=['GET'])
def exportResultToPDF(id,templateid):
	try:
		exeResult=DB.getExeResult(exeId=id)
		exeOb=DB.getExeOBHist(projectId=projectSession(),exeId=id)
		name=exeOb[1] + " (" + exeOb[2] + ")"
		template=UP.getTemplateFile(fileId=templateid)
		path='pdfTemplates/'+str(template[2])
		html=render_template(path,exePDF=exeResult,exeOb=exeOb)
		filename="./downloads/"+name.replace(".","_").replace("/","-")+".pdf"
		DB.insertFile(name=filename,projectId=projectSession())
	except Exception as e:
		exc_type, exc_obj, exc_tb = sys.exc_info()
		fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in exportResultToPDF method in
			""")
		print(exc_type, fname, exc_tb.tb_lineno)
		return False
	
	path=bytes(r'./wkhtmltopdf/bin/wkhtmltopdf.exe','utf-8')
	config = pdfkit.configuration(wkhtmltopdf=path)
	pdfkit.from_string(html, filename, configuration=config)
	
	return filename

@app.route('/exportResultToXLSX/<int:id>', methods=['GET'])
def exportResultToXLSX(id):
	try:
		exeResult=DB.getExeResult(exeId=id)
		exeOb=DB.getExeOBHist(projectId=projectSession(),exeId=id)
		name=exeOb[1] + " (" + exeOb[2] + ")"
		filename="./downloads/"+name.replace(".","_").replace("/","-")+".xlsx"
		DB.insertFile(name=filename,projectId=projectSession())
	except:
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in exportCaseToXLSX method
			""")
	try:
		workbook = xlsxwriter.Workbook(filename)
		
		green_bg = workbook.add_format()
		green_bg.set_pattern(1)  # This is optional when using a solid fill.
		green_bg.set_bg_color('#d0e9c6')
		
		blue_bg = workbook.add_format()
		blue_bg.set_pattern(1)  # This is optional when using a solid fill.
		blue_bg.set_bg_color('#c4e3f3')
		
		red_bg = workbook.add_format()
		red_bg.set_pattern(1)  # This is optional when using a solid fill.
		red_bg.set_bg_color('#f2dede')
		
		yellow_bg = workbook.add_format()
		yellow_bg.set_pattern(1)  # This is optional when using a solid fill.
		yellow_bg.set_bg_color('#faf2cc')
		
		worksheet = workbook.add_worksheet()
		bold = workbook.add_format({'bold': True})
		worksheet.set_column('A:D', len(exeResult)+4)
		worksheet.write('A1', 'Title:', bold)
		worksheet.write('A2', 'Description:', bold)
		worksheet.write('B1', name)
		worksheet.write('B2', "")
		worksheet.write('A4', 'Case Name', bold)
		worksheet.write('B4', 'Result', bold)
		worksheet.write('C4', 'Assigned To', bold)
		worksheet.write('D4', 'Duration', bold)
		for id,item in enumerate(exeResult):
			if item[2] == "RUN":
				style = green_bg
			if item[2] == "FAILED":
				style = red_bg
			if item[2] == "NOTRUN" or item[2] == "SKIPPED":
				style = yellow_bg
			if item[2] == "NOTIMP":
				style = blue_bg
			worksheet.write(id+4, 0, render_template("caseToDoc.html", caseDoc=item[1]), style)
			worksheet.write(id+4, 1, render_template("caseToDoc.html", caseDoc=item[2]), style)
			worksheet.write(id+4, 2, "", style)
			worksheet.write(id+4, 3, "", style)
		workbook.close()
	except Exception as e:
		exc_type, exc_obj, exc_tb = sys.exc_info()
		fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
		print("Unexpected error: Unexcepted error occured during the worksheet method in exportCaseToXLSX")
		print(exc_type, fname, exc_tb.tb_lineno)
	return filename	

@app.route('/advancedReportExportXLSX', methods=['POST'])
def advancedReportExportXLSX():
	try:
		request.form['filename']
		exes=[]
		for k in request.form.getlist('exeIds[]'):
			exes.append(list(DB.getReportGenRes(exeId=k)))
		temp=[]
		for exe in exes:
			tempExe=[]
			for step in exe:
				vars=[]
				for var in request.form.getlist('variableBox'):
					v=DB.getGenResVar(varId=var,exeId=step[0],stepId=step[8])
					if v != None:
						vars.append(v)
					else:
						vars.append("")
				tags=[]
				for area in request.form.getlist('areaBox'):
					a=DB.getCaseAreaWithArea(caseId=step[5],areaId=area)
					if a != None:
						tags.append(a)
					else:
						tags.append("")
				#Tstep=[step,tags,vars]
				tempExe.append([step,tags,vars])
			temp.append(tempExe)
		header = ['Object', 'Version', 'Execution', 'Case', 'Case Result', 'Step Action', 'Step Result']
		for area in request.form.getlist('areaBox'):
			header.append(DB.getArea(areaId=area)[1])
		for var in request.form.getlist('variableBox'):
			header.append(DB.getVariable(varId=var)[2])
		filename="./downloads/"+request.form['filename'].replace(".","_").replace("/","-")+".xlsx"
		DB.insertFile(name=filename,projectId=projectSession())
	except Exception as e:
		exc_type, exc_obj, exc_tb = sys.exc_info()
		fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
		print("""
			Unexpected error: Unexcepted error occured during the dataQuery in advancedReportExportXLSX method in
			""")
		print(exc_type, fname, exc_tb.tb_lineno)
		return False
	try:
		workbook = xlsxwriter.Workbook(filename)
		worksheet = workbook.add_worksheet()
		for it,element in enumerate(header):
			worksheet.write(0, it, element)
		it1=0
		for exe in temp:
			for steps in exe:
				it2=0
				worksheet.write(it1+1,it2,steps[0][3])
				worksheet.write(it1+1,it2+1,steps[0][4])
				worksheet.write(it1+1,it2+2,steps[0][1])
				worksheet.write(it1+1,it2+3,steps[0][7])
				worksheet.write(it1+1,it2+4,steps[0][6])
				worksheet.write(it1+1,it2+5,render_template("caseToDoc.html", caseDoc=steps[0][11]))
				worksheet.write(it1+1,it2+6,steps[0][9])
				print(it2)
				it2=7
				for t in steps[1]:
					if t != "":
						if t[2] != None:
							worksheet.write(it1+1,it2,t[2])
						else:
							worksheet.write(it1+1,it2,"")
					else:
						worksheet.write(it1+1,it2,"")
					it2=it2+1
				for v in steps[2]:
					if v != "":
						if v[2] != None:
							worksheet.write(it1+1,it2,v[2])
						else:
							worksheet.write(it1+1,it2,"")
					else:
						worksheet.write(it1+1,it2,"")
					it2=it2+1
				it1=it1+1
		workbook.close()
		return filename
	except Exception as e:
		exc_type, exc_obj, exc_tb = sys.exc_info()
		fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
		print("""
			Unexpected error: Unexcepted error occured during the worksheet method in advancedReportExportXLSX
			""")
		print(exc_type, fname, exc_tb.tb_lineno)
		return False
	return "error"
	
@app.route('/tableSchema/<name>', methods=['GET'])
def tableSchema(name):
	schema=DB.getSchema(name=name);
	filename="./downloads/"+name+".txt"
	DB.insertFile(name=filename,projectId=projectSession())
	text_file = open(filename, "w")
	text_file.write(schema[0])
	text_file.close()
	return filename
	
@app.route('/tableData/<name>', methods=['GET'])
def tableData(name):
	data=DB.getDataFromTable(name=name);
	filename="./downloads/"+name.replace(".","_").replace("/","-")+".xlsx"
	DB.insertFile(name=filename,projectId=projectSession())
	workbook = xlsxwriter.Workbook(filename)
	worksheet = workbook.add_worksheet()
	for id,item in enumerate(data):
		for l,k in enumerate(item):
			worksheet.write(id, l-1, k)
		workbook.close()
	return filename
	
@app.route('/deleteFileDownload/<int:id>', methods=['GET'])
def deleteFileDownload(id):
	path=DB.deleteFile(id=id)
	os.remove(path)
	files=DB.getDownloadFiles(projectId=projectSession())
	return render_template("exportimport.html", newfiles="true",filesNew=files)
	
@app.route('/deleteAllFilesDownload', methods=['GET'])
def deleteAllFilesDownload():
	files=DB.getDownloadFiles(projectId=projectSession())
	for k in files:
		newfile=deleteFileDownload(k[0])
	return render_template("exportimport.html", newfiles="true",filesNew=newfile)

@app.route('/refreshDownloadFiles', methods=['GET'])
def refreshDownloadFiles():
	files=DB.getDownloadFiles(projectId=projectSession())
	return render_template("exportimport.html", newfiles="true",filesNew=file)
	
@app.route('/deleteAllTemplates', methods=['GET'])
def deleteAllTemplates():
	files=UP.deleteAllTemplates()
	for k in files:
		os.remove(k[1])
	temps=UP.getTemplates()
	return render_template("exportimport.html", templatesFile=temps)

#----- Email ----
@app.route('/saveEmail', methods=['POST'])
def SaveEmail():
	result = DB.SaveEmail(name=request.form['name'],userId=request.form['user'],D=request.form['D'],W=request.form['W'],U=request.form['U'],projectId=projectSession())
	result = DB.getEmails(projectId=projectSession())
	user = DB.getUsers()
	return render_template("exportimport.html", emailtable=result, user=user)
	
@app.route('/editEmail/<int:id>', methods=['POST'])
def EditEmail(id):
	DB.EditEmail(id=id,name=request.form['name'],userId=request.form['user'],D=request.form['D'],W=request.form['W'],U=request.form['U'])
	result = DB.getEmails(projectId=projectSession())
	user = DB.getUsers()
	return render_template("exportimport.html", emailtable=result, user=user)
	
@app.route('/deleteEmail/<int:id>', methods=['GET'])
def DeleteEmail(id):
	DB.DeleteEmail(id=id)
	return "OK"
	
#----- Email ----
@app.route('/updateEmailConfig', methods=['POST'])
def UpdateEmailConfig():
	DB.UpdateEmailConfig(email=request.form['email'],smtp=request.form['smtp'], \
			port=request.form['port'],projectId=projectSession())
	return "Ok"
	
@app.route('/sendEmail/<type>/<obid>', methods=['GET'])
def SendEmail(type,obid):
	if obid == 'latest':
		obid = DB.getLatestObject()
	if type == 'D':
		result = DB.getDailyReportToSend(obid=obid,projectId=projectSession())
		html = render_template("sendDaily.html",daily=result)
		address = DB.getEmailAddress(D=1,W=0,U=0,projectId=projectSession())
		config = DB.GetEmailConfig(projectId=projectSession())
		return html
		#Messenger.EmailSender(config[1],config[2],config[3],address,html)
	if type == 'W':
		result = DB.getWeeklyReportToSend(projectId=projectSession())
		address = DB.getEmailAddress(D=0,W=1,U=0,projectId=projectSession())
		config = DB.GetEmailConfig(projectId=projectSession())
		html = render_template("sendWeekly.html",weekly=result)
		Messenger.EmailSender(config[1],config[2],config[3],address,html)
	return "Ok"
	
#-----Unit Test Input----
@app.route('/GetAreas/<int:id>', methods=['GET'])
def GetAreas(id):
	return json.dumps(DB.getAreas(projectId=id))

@app.route('/ObjectInput', methods=['POST'])	
def ObjectInput():
	print(ast.literal_eval(request.data.decode("utf-8")))
	data = ast.literal_eval(request.data.decode("utf-8"))
	print(data["name"])
	ID=DB.save_object(name=data["name"],hardware=data["hardware"],desc=data["desc"],
		version=data["version"],projectId=data["projectId"],areas=data['areaBox'])
	return json.dumps(ID)

@app.route('/UnitGetSet', methods=['POST'])	
def UnitGetSet():
	data = ast.literal_eval(request.data.decode("utf-8"))
	return json.dumps(DB.getSetCases(id=
			DB.getSetByName(sets=data['sets'],projectId=data['projectId']
			,active=1,update=0)[0][0]))

@app.route('/UnitUser', methods=['POST'])	
def UnitUser():
	data = ast.literal_eval(request.data.decode("utf-8"))
	return json.dumps(DB.getUSerByName(name=data['name']))
	
@app.route('/UnitNewExe', methods=['POST'])
def UnitNewExe():
	data = ast.literal_eval(request.data.decode("utf-8"))
	print(data)
	exeId=DB.saveExe(name=data["title"],testObject=data["TO"],projectId=data['projectId'],
					areas=data['areaBox'],userId=data["userId"],
					dynamic=data['dynamicArea'])
	DB.saveCaseExe(ID=data['ID'],exeID=exeId)
	return json.dumps(exeId)

@app.route('/UnitCaseExeId', methods=['POST'])
def UnitCaseExeId():
	data = ast.literal_eval(request.data.decode("utf-8"))
	caseExeId=DB.unitGetCaseExeId(exeId=data["exeId"],caseName=data["caseName"])
	return json.dumps(caseExeId)
	
@app.route('/UnitStepIdByCaseName', methods=['POST'])
def UnitStepIdByCaseName():
	data = ast.literal_eval(request.data.decode("utf-8"))
	stepIds=DB.UnitStepIdByCaseName(exeId=data["exeId"],caseName=data["caseName"])
	return json.dumps(stepIds)
	
@app.route('/UnitCaseIdByCaseName', methods=['POST'])
def UnitCaseIdByCaseName():
	data = ast.literal_eval(request.data.decode("utf-8"))
	stepIds=DB.UnitCaseIdByCaseName(exeId=data["exeId"],caseName=data["caseName"])
	return json.dumps(stepIds)
	
# set the secret key.  keep this really secret:
app.secret_key = os.urandom(24)
		
if __name__ == "__main__":
	app.debug = True
	#toolbar = DebugToolbarExtension(app)
	app.run(host= '0.0.0.0')
	#app.run()